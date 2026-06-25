import sys
import re
import os
import subprocess
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor, Inches

def make_direct_report():
    input_path = Path("/home/niranjan/Desktop/NEXUSDESK_STATUS_REPORT.md")
    output_docx = Path("/home/niranjan/Desktop/NexusDesk_Project_Status_Report.docx")
    
    if not input_path.exists():
        print(f"Error: {input_path} not found.")
        sys.exit(1)
        
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    PRIMARY_COLOR = (45, 45, 45) # Dark Ink
    TEXT_COLOR = (60, 60, 60)
    ACCENT_COLOR = (184, 135, 42) # Amber Accent
    
    def set_font(run, font_name="Arial", size_pt=11, color_rgb=TEXT_COLOR, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(*color_rgb)
        run.bold = bold
        run.italic = italic

    # Add Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run("NexusDesk — Complete Feature Specification")
    set_font(run, font_name="Arial", size_pt=20, color_rgb=PRIMARY_COLOR, bold=True)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run("Generated directly from NEXUSDESK_STATUS_REPORT.md")
    set_font(sub_run, size_pt=10, color_rgb=(120, 120, 120), italic=True)
    
    content = input_path.read_text(encoding="utf-8")
    lines = content.split("\n")
    
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("---") or stripped.startswith("# "):
            continue
            
        if stripped.startswith("## "):
            heading_text = stripped.replace("## ", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(heading_text)
            set_font(run, size_pt=14, color_rgb=ACCENT_COLOR, bold=True)
            
        elif stripped.startswith("* "):
            bullet_text = stripped.replace("* ", "").strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            # Parse inline bolding **text**
            parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    bold_text = part[2:-2]
                    run = p.add_run(bold_text)
                    set_font(run, size_pt=11, color_rgb=TEXT_COLOR, bold=True)
                else:
                    run = p.add_run(part)
                    set_font(run, size_pt=11, color_rgb=TEXT_COLOR)
                    
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    bold_text = part[2:-2]
                    run = p.add_run(bold_text)
                    set_font(run, size_pt=11, color_rgb=TEXT_COLOR, bold=True)
                else:
                    run = p.add_run(part)
                    set_font(run, size_pt=11, color_rgb=TEXT_COLOR)
                    
    doc.save(str(output_docx))
    print("✅ Word document successfully generated directly (no paragraphs)!")
    
    # Try converting to PDF using LibreOffice
    try:
        print("📄 Converting to PDF using LibreOffice...")
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf", 
            "--outdir", "/home/niranjan/Desktop", str(output_docx)
        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
        print("✅ PDF successfully generated directly!")
    except Exception as e:
        print(f"⚠️ PDF conversion skipped or failed: {e}")

if __name__ == "__main__":
    make_direct_report()
